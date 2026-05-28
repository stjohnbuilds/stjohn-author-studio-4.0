import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt


def configure_document(document):
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    heading = document.styles["Heading 1"]
    heading.font.name = "Times New Roman"
    heading.font.size = Pt(18)
    heading.font.bold = True


def add_chapter(document, chapter, is_first):
    if not is_first:
        document.add_page_break()

    heading = document.add_heading(chapter["title"], level=1)
    heading.paragraph_format.space_after = Pt(18)

    for text in chapter.get("paragraphs", []):
        paragraph = document.add_paragraph(text)
        paragraph.paragraph_format.first_line_indent = Inches(0.25)
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.paragraph_format.space_after = Pt(10)


def main():
    if len(sys.argv) != 3:
        raise SystemExit("Usage: build-manuscript-docx.py project.json output.docx")

    project_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    project = json.loads(project_path.read_text())

    document = Document()
    configure_document(document)
    document.core_properties.title = project.get("title", "Quill & Ink InDesign Export Test")
    document.core_properties.author = "StJohn Author Studio test pack"

    for index, chapter in enumerate(project.get("chapters", [])):
        add_chapter(document, chapter, index == 0)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


if __name__ == "__main__":
    main()

